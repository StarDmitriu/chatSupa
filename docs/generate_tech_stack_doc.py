#!/usr/bin/env python3
"""Генерирует TECH_STACK_AND_CAPACITY.docx из встроенного текста (без внешних секретов)."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    t = doc.add_paragraph()
    t.add_run("Технический стек и ёмкость сервера").bold = True
    t.runs[0].font.size = Pt(18)
    doc.add_paragraph(
        "Проект: платформа массовых рассылок (WhatsApp + Telegram), личный кабинет, "
        "шаблоны, кампании, оплаты. Документ сформирован по состоянию репозитория /var/www."
    )
    doc.add_paragraph(
        "Формат: .docx (Microsoft Word). Устаревший бинарный .doc на сервере не генерировался; "
        "при необходимости откройте этот файл в Word и «Сохранить как» → .doc."
    )

    add_heading(doc, "1. Архитектура (высокий уровень)", 1)
    add_bullets(
        doc,
        [
            "Клиент: браузер → (опционально) nginx TLS → Next.js standalone на 127.0.0.1:3001.",
            "Next.js проксирует HTTP-запросы с префикса /api/* на бэкенд NestJS (по умолчанию localhost:3000; в Docker — BACKEND_INTERNAL_URL).",
            "Бэкенд: один процесс Node (NestJS), внутри — REST API, интеграции WA/TG, планировщик повторов кампаний, BullMQ-воркер отправки.",
            "Очередь: Redis + BullMQ, очередь campaign-send.",
            "Данные: Supabase (PostgreSQL + API), service role на бэкенде для серверных операций.",
            "Процесс-менеджер: PM2 (backend + frontend), автоперезапуск, лимиты памяти.",
        ],
    )

    add_heading(doc, "2. Фронтенд", 1)
    add_bullets(
        doc,
        [
            "Фреймворк: Next.js 16.x (App Router), output: «standalone» для одного артефакта деплоя.",
            "React 19.2, TypeScript 5.",
            "UI: Ant Design (^6.1).",
            "Данные/кэш клиента: SWR.",
            "Прочее: js-cookie, qrcode.react.",
            "Сборка: NODE_OPTIONS --max-old-space-size=6144 (тяжёлый билд); postbuild копирует static и public в standalone.",
            "Порт production: 3001; скрипт start: node .next/standalone/server.js.",
            "Тесты: Vitest; E2E: Playwright.",
            "Линтинг: ESLint 9 + eslint-config-next.",
            "Прокси API: next.config.ts rewrites /api → бэкенд.",
            "src/proxy.ts: защита префиксов /dashboard, /cabinet, /admin; Cache-Control no-store для HTML/RSC после деплоя (избежание ChunkLoadError).",
        ],
    )

    add_heading(doc, "3. Бэкенд", 1)
    add_bullets(
        doc,
        [
            "Платформа: NestJS 11, Node.js, TypeScript 5.7.",
            "HTTP: @nestjs/platform-express, глобальный ValidationPipe (whitelist, forbidNonWhitelisted, transform).",
            "Аутентификация: @nestjs/jwt, @nestjs/passport, passport-jwt; CORS настраивается через CORS_ORIGINS.",
            "База/API: @supabase/supabase-js; миграции SQL в каталоге supabase/migrations (часть схемы в облаке Supabase).",
            "Очередь: bullmq + ioredis; поддержка REDIS_URL или REDIS_HOST/PORT/PASSWORD/DB; TLS для rediss:.",
            "WhatsApp: @whiskeysockets/baileys (WebSocket-сессии, QR, отправка в группы).",
            "Telegram: пакет telegram (GramJS, MTProto), интеграция с пользовательскими сессиями.",
            "Время/расписание: luxon; часовой пояс по умолчанию DEFAULT_TZ=Europe/Moscow.",
            "Логи: pino (в зависимостях).",
            "Файлы/импорт: multer, papaparse (CSV), axios.",
            "Ошибки HTTP-уровня (WA): @hapi/boom.",
            "Платежи: модуль Prodamus (переменные PRODAMUS_* в .env.example).",
            "SMS: SMS.RU (опционально SMSRU_*).",
            "Таблицы: Google Apps Script (опционально APPS_SCRIPT_*).",
            "Админка: ADMIN_PANEL_PASSWORD (опционально).",
            "Повтор кампаний: CAMPAIGN_REPEAT_ENABLED, CAMPAIGN_REPEAT_TICK_MS (см. CampaignRepeatService).",
        ],
    )

    add_heading(doc, "4. Функциональные модули (Nest)", 1)
    add_bullets(
        doc,
        [
            "AuthModule — вход, JWT.",
            "SupabaseModule — клиент к БД.",
            "WhatsappModule — сессии и отправка WA.",
            "TelegramModule — TG.",
            "TemplatesModule — шаблоны и цели (targets).",
            "CampaignsModule — кампании, волны, jobs, ресинк расписания, повторы.",
            "QueueModule — BullMQ очередь и воркер отправки.",
            "SheetsModule — синхронизация с таблицами.",
            "AdminModule, PaymentsModule, LeadsModule.",
            "SubscriptionsService — проверка доступа по каналу при старте/повторе волн.",
        ],
    )

    add_heading(doc, "5. Очередь и отправка", 1)
    add_bullets(
        doc,
        [
            "Имя очереди BullMQ: campaign-send.",
            "Воркер CampaignBullWorker: concurrency = 1 — одновременно обрабатывается одна задача отправки на весь инстанс (глобальная сериализация).",
            "Сохранение ритма: интервалы из campaign_jobs.scheduled_at при «догоне» очереди; сброс ритма после простоя и при больших скачках расписания.",
            "Адаптивные задержки при TG FLOOD (обучение по частоте за 5 минут).",
            "Повтор волны: только если нет pending/processing jobs с scheduled_at ≤ now; иначе repeat откладывается.",
        ],
    )

    add_heading(doc, "6. Среда выполнения Node.js", 1)
    add_bullets(
        doc,
        [
            "На продакшен-VPS: Node.js v20.x (LTS), согласовано с Dockerfile backend/frontend (node:20-alpine).",
        ],
    )

    add_heading(doc, "7. Инфраструктура деплоя", 1)
    add_bullets(
        doc,
        [
            "Вариант A — PM2 на хосте (ecosystem.config.cjs): backend — cwd backend, node dist/main.js, instances: 1, max_memory_restart: 1024M, PORT 3000, BIND_HOST 127.0.0.1.",
            "PM2: frontend — .next/standalone/server.js, instances: 1, max_memory_restart: 500M, PORT 3001, HOSTNAME 127.0.0.1.",
            "NEXT_SERVER_ACTIONS_ENCRYPTION_KEY — задавать через env/PM2 для стабильности Server Actions после деплоя (не хранить в репозитории).",
            "Пример nginx: deploy/nginx-frontend-next.example.conf — прокси на :3001, кэш для /_next/static.",
            "Redis на хосте: обычно 127.0.0.1:6379 (проверка: redis-cli ping).",
            "Домен в CORS по умолчанию: chatrassylka.ru / www (см. main.ts).",
        ],
    )

    add_heading(doc, "8. Альтернатива: Docker Compose", 1)
    add_bullets(
        doc,
        [
            "docker-compose.yml: сервис redis (redis:7-alpine, AOF, лимиты cpu/mem, без публикации порта наружу).",
            "backend: сборка из ./backend, порт 127.0.0.1:3000, REDIS_URL=redis://redis:6379, CAMPAIGN_REPEAT_ENABLED=true, healthcheck wget.",
            "frontend: сборка из ./frontend, NEXT_PUBLIC_BACKEND_URL=/api, BACKEND_INTERNAL_URL=http://backend:3000, порт 127.0.0.1:3001, shm_size 256mb (важно для Next).",
            "nginx: профиль docker-nginx, образ nginx:alpine, тома для deploy/nginx/default.conf и certbot (Let’s Encrypt).",
            "docker-compose.pm2.yml — вариант композиции с PM2 (см. файл в репозитории).",
            "В Docker лимиты: backend mem ~1536m, frontend ~768m, redis 512m — ориентир для планирования RAM хоста (сумма + ОС).",
        ],
    )

    add_heading(doc, "9. Переменные окружения (обзор)", 1)
    add_bullets(
        doc,
        [
            "Backend (.env.example): PORT, BIND_HOST, CORS_ORIGINS, SUPABASE_*, JWT_SECRET, REDIS_*, DEFAULT_TZ, Prodamus, опционально TG_API_ID/HASH, SMS, OTP_*, CAMPAIGN_REPEAT_*, ADMIN_PANEL_PASSWORD, SUPABASE_DB_URL для миграций.",
            "Frontend: NEXT_PUBLIC_BACKEND_URL (часто /api при прокси Next).",
        ],
    )

    add_heading(doc, "10. Текущий сервер (снимок)", 1)
    add_bullets(
        doc,
        [
            "ОС: Ubuntu 22.04.5 LTS, ядро 5.15.x.",
            "CPU: 2 логических ядра (nproc).",
            "RAM: ~5.8 GiB всего; доступно под нагрузку заметно меньше из‑за кэша и уже запущенных процессов.",
            "Swap: отсутствует (0) — при пике памяти возможен OOM-killer.",
            "Диск: ~20 GiB корень, занято ~77% — запас места ограничен; рост логов и билдов быстро забивает диск.",
            "Сеть: типичный VPS (детали провайдера в панели; в документе не фиксировались).",
        ],
    )

    add_heading(doc, "11. Ёмкость и «сколько выдержит»", 1)
    doc.add_paragraph(
        "Точное число «пользователей» без нагрузочного теста не существует. Ниже — оценка по узким местам архитектуры."
    )
    add_heading(doc, "11.1. Конечные получатели сообщений (чаты/группы)", 2)
    add_bullets(
        doc,
        [
            "Лимит не в «количестве людей в БД», а в скорости отправки: задержки между группами и шаблонами, лимиты WhatsApp/Telegram, flood-wait.",
            "Один воркер BullMQ (concurrency 1) упорядочивает нагрузку, но не даёт параллельно слать несколько сообщений из очереди.",
            "Практическая пропускная способность: от минут до часов на большую волну (сотни/тысячи jobs) — см. логи WAVE CREATED и интервалы шаблонов.",
        ],
    )
    add_heading(doc, "11.2. Клиенты платформы (владельцы кабинетов)", 2)
    add_bullets(
        doc,
        [
            "Веб: Next standalone + 500M лимит PM2 — для типичных страниц кабинета десятки одновременных сессий обычно комфортны; сотни — возможны при лёгких страницах и без тяжёлых запросов, но без профилирования гарантий нет.",
            "API: один инстанс Nest с лимитом 1024M — основная память уходит на Baileys/GramJS сессии активных пользователей; много одновременных WA/TG-подключений быстро потребляет RAM.",
            "Supabase: лимиты тарифа (connections, egress, rate limits) могут стать потолком раньше CPU сервера.",
        ],
    )
    add_heading(doc, "11.3. Ориентиры (очень грубо)", 2)
    add_bullets(
        doc,
        [
            "«До N человек одновременно в вебе»: порядок 30–100 как разумный диапазон для оценки без нагрузочного теста на данном VPS; выше — нужны замеры.",
            "«Сколько получателей в сутки»: определяется лимитами мессенджеров и паузами в шаблонах, а не только сервером.",
        ],
    )

    add_heading(doc, "12. Что улучшить (приоритеты)", 1)
    add_bullets(
        doc,
        [
            "Диск: очистка логов PM2, старых билдов, мониторинг df; при росте — увеличить том или ротация логов.",
            "Память: добавить swap 2–4 GiB или увеличить RAM VPS; пересмотреть max_memory_restart при стабильном RSS.",
            "Наблюдаемость: централизованные логи, метрики (CPU/RAM/Redis), алерты на OOM и 5xx.",
            "Очередь: вынести воркер отправки в отдельный процесс/контейнер; при необходимости — осторожно увеличить concurrency с учётом лимитов WA/TG (риск банов).",
            "Масштабирование API: горизонтальное масштабирование Nest затруднено из‑за состояния Baileys в памяти — обычно нужны sticky sessions или выделенные воркеры только для очереди, а WA-сессии на выделенных нодах.",
            "База: индексы под частые запросы campaign_jobs; при росте — пул соединений и тариф Supabase.",
            "Redis: отдельный managed Redis при отказоустойчивости; бэкапы AOF/RDB.",
            "Безопасность: секреты только в env; не коммитить ключи; регулярные npm audit / обновления зависимостей.",
            "CDN/кэш для статики Next уже заложен в nginx-примере; для глобальной аудитории — CDN перед origin.",
        ],
    )

    add_heading(doc, "13. Тестирование и качество", 1)
    add_bullets(
        doc,
        [
            "Backend: Jest.",
            "Frontend: Vitest, Playwright E2E.",
            "ESLint на обоих слоях.",
        ],
    )

    add_heading(doc, "14. База данных и миграции", 1)
    add_bullets(
        doc,
        [
            "Основное хранилище: PostgreSQL через Supabase.",
            "SQL-миграции в репозитории: supabase/migrations/ (пример: wa_phone для групп, индексы/страницы для telegram_groups).",
            "RLS и прочие политики — по проекту в Supabase Dashboard / отдельные скрипты (npm run migrate:rls в backend).",
        ],
    )

    add_heading(doc, "15. Репозиторий и скрипты", 1)
    add_bullets(
        doc,
        [
            "Корень: ecosystem.config.cjs, deploy/, scripts/, supabase/migrations/.",
            "Документация: DEPLOY_PM2.md (упоминается в ecosystem), docs/PAYMENT_DEBUG.md и др.",
        ],
    )

    doc.add_paragraph("")
    p = doc.add_paragraph(
        "Документ сгенерирован скриптом docs/generate_tech_stack_doc.py; при изменении стека обновите скрипт или правьте .docx вручную."
    )
    p.runs[0].italic = True

    out = "/var/www/docs/TECH_STACK_AND_CAPACITY.docx"
    doc.save(out)
    print("Wrote", out)

if __name__ == "__main__":
    main()
